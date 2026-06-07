"""
DACL Document Extractor — PDF / TXT / Excel / DOCX / CSV → plain text.

Two use-cases:
  1. Policy compilation:  PDF with 100K+ business rules → text → LLM compiler
  2. Runtime fact injection: PDF case document (insurance claim, shipment waybill, etc.)
                             → text → LLM fact extractor → DACL Rete engine

Supported formats:
  - PDF  (.pdf)           via pypdf
  - Excel (.xlsx, .xls)   via openpyxl
  - Word  (.docx)         via python-docx
  - CSV   (.csv)          built-in
  - Plain text (.txt, .md, .xml, etc.)
"""
from __future__ import annotations
import csv
from io import BytesIO, StringIO

# How many chars to send to the LLM (compile = more, fact extract = less)
MAX_CHARS_POLICY = 200_000   # large rule documents (1000+ rules) — GPT-4o 128K ctx = ~200K chars
MAX_CHARS_FACTS  =  12_000   # case documents — only facts matter

SUPPORTED_EXTENSIONS = [".pdf", ".txt", ".csv", ".xlsx", ".xls", ".docx", ".md", ".xml"]


def extract_text(
    content: bytes,
    filename: str,
    content_type: str = "",
    max_chars: int = MAX_CHARS_POLICY,
) -> str:
    """
    Extract plain text from uploaded file bytes.

    Args:
        content:      Raw file bytes.
        filename:     Original filename (used to detect format).
        content_type: MIME type from the upload (fallback detection).
        max_chars:    Truncate output to this many chars (default: policy mode).

    Returns:
        Plain text string, truncated to max_chars if necessary.

    Raises:
        RuntimeError: if extraction fails or a required library is not installed.
    """
    fname = (filename or "").lower()

    if fname.endswith(".pdf") or content_type in ("application/pdf", "application/x-pdf"):
        text = _extract_pdf(content)
    elif fname.endswith((".xlsx", ".xls")) or content_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        text = _extract_excel(content)
    elif fname.endswith(".docx") or content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = _extract_docx(content)
    elif fname.endswith(".csv") or content_type == "text/csv":
        text = _extract_csv(content)
    else:
        # Plain text, Markdown, XML, etc.
        text = content.decode("utf-8", errors="replace")

    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[... document truncated at {max_chars:,} characters ...]"

    return text.strip()


def _extract_pdf(content: bytes) -> str:
    """Extract text from all pages of a PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError(
            "pypdf is required for PDF uploads.  "
            "Install it with:  pip install pypdf"
        )

    try:
        reader = PdfReader(BytesIO(content))
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"--- Page {i + 1} ---\n{page_text}")
        return "\n\n".join(pages)
    except Exception as exc:
        raise RuntimeError(f"PDF extraction failed: {exc}") from exc


def _extract_excel(content: bytes) -> str:
    """Extract text from all sheets of an Excel workbook using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise RuntimeError(
            "openpyxl is required for Excel uploads.  "
            "Install it with:  pip install openpyxl"
        )

    try:
        wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
        sheets: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):  # skip fully empty rows
                    rows.append(" | ".join(cells))
            if rows:
                sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets)
    except Exception as exc:
        raise RuntimeError(f"Excel extraction failed: {exc}") from exc


def _extract_docx(content: bytes) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX uploads.  "
            "Install it with:  pip install python-docx"
        )

    try:
        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        raise RuntimeError(f"DOCX extraction failed: {exc}") from exc


def _extract_csv(content: bytes) -> str:
    """Extract text from a CSV file, converting rows to pipe-delimited text."""
    try:
        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(StringIO(text))
        rows: list[str] = []
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append(" | ".join(row))
        return "\n".join(rows)
    except Exception as exc:
        raise RuntimeError(f"CSV extraction failed: {exc}") from exc
